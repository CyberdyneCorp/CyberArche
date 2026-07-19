"""document-import spec: the FileExtractor adapter.

Markdown now yields structured blocks (not flat paragraphs); Word documents are
mapped by paragraph style with tables; unreadable Word input is rejected.
"""

from __future__ import annotations

import io

import pytest
from docx import Document as Docx

from cyberarche.adapters.outbound.extraction.files import FileExtractor
from cyberarche.domain.errors import ValidationFailed


def _docx_bytes(build) -> bytes:
    document = Docx()
    build(document)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_markdown_yields_structured_blocks_not_flat_paragraphs():
    extractor = FileExtractor()
    text = "# Title\n\nIntro.\n\n- one\n- two\n\n```py\nx = 1\n```"
    blocks = extractor.extract_blocks(filename="notes.md", content=text.encode())
    types = [b["type"] for b in blocks]
    assert types == ["heading", "paragraph", "bulleted_list", "bulleted_list", "code"]
    assert blocks[0]["data"] == {"text": "Title", "level": 1}
    assert all(b["id"] for b in blocks)


def test_markdown_extension_variant_is_also_structured():
    extractor = FileExtractor()
    blocks = extractor.extract_blocks(filename="a.markdown", content=b"## Sub")
    assert blocks[0]["type"] == "heading"
    assert blocks[0]["data"]["level"] == 2


def test_plain_text_stays_flat_paragraphs():
    extractor = FileExtractor()
    blocks = extractor.extract_blocks(filename="a.txt", content=b"# not a heading")
    assert [b["type"] for b in blocks] == ["paragraph"]
    assert blocks[0]["data"]["text"] == "# not a heading"


def test_docx_maps_headings_lists_quote_and_paragraphs_by_style():
    def build(document):
        document.add_heading("Report", level=1)
        document.add_heading("Details", level=2)
        document.add_paragraph("Body text.")
        document.add_paragraph("bullet item", style="List Bullet")
        document.add_paragraph("numbered item", style="List Number")
        document.add_paragraph("a quotation", style="Quote")
        document.add_paragraph("   ")  # empty -> skipped

    extractor = FileExtractor()
    blocks = extractor.extract_blocks(filename="r.docx", content=_docx_bytes(build))
    types = [b["type"] for b in blocks]
    assert types == [
        "heading",
        "heading",
        "paragraph",
        "bulleted_list",
        "numbered_list",
        "quote",
    ]
    assert blocks[0]["data"] == {"text": "Report", "level": 1}
    assert blocks[1]["data"]["level"] == 2
    assert blocks[3]["data"]["text"] == "bullet item"


def test_docx_tables_become_table_blocks():
    def build(document):
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Age"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "30"

    extractor = FileExtractor()
    blocks = extractor.extract_blocks(filename="t.docx", content=_docx_bytes(build))
    table_blocks = [b for b in blocks if b["type"] == "table"]
    assert len(table_blocks) == 1
    data = table_blocks[0]["data"]
    assert data["header"] == ["Name", "Age"]
    assert data["rows"] == [["Alice", "30"]]
    assert data["source"] == "docx"


def test_unreadable_docx_raises_validation_failed():
    extractor = FileExtractor()
    with pytest.raises(ValidationFailed):
        extractor.extract_blocks(filename="broken.docx", content=b"not a real docx")
